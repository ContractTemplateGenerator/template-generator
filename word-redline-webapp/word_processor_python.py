#!/usr/bin/env python3
"""
Python-based Word document processor using python-docx
Provides advanced Word document manipulation capabilities similar to the MCP server
"""

import sys
import json
import argparse
from pathlib import Path
import re
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX
from datetime import datetime

def create_track_changes_element(doc, change_type, author="Redline Processor"):
    """Create track changes XML elements."""
    change_id = str(hash(datetime.now().isoformat()) % 1000000)
    date_str = datetime.now().isoformat()
    
    if change_type == "delete":
        element = OxmlElement('w:del')
    elif change_type == "insert":
        element = OxmlElement('w:ins')
    else:
        raise ValueError(f"Unsupported change type: {change_type}")
    
    element.set(qn('w:id'), change_id)
    element.set(qn('w:author'), author)
    element.set(qn('w:date'), date_str)
    
    return element

def apply_redline_instruction(doc, instruction):
    """Apply a single redline instruction to the document."""
    instruction_type = instruction.get('type')
    original = instruction.get('original', '')
    
    print(f"Processing: {original}")
    
    if instruction_type == 'replace':
        find_text = instruction['find']
        replace_text = instruction['replace']
        
        print(f"Looking for: '{find_text}' to replace with: '{replace_text}'")
        
        # Find and replace text while adding track changes
        found = False
        for paragraph in doc.paragraphs:
            if find_text in paragraph.text:
                print(f"Found in paragraph: '{paragraph.text[:100]}...'")
                
                # Create new paragraph with track changes
                # This is a simplified approach - for full track changes, we'd need to manipulate XML directly
                paragraph.text = paragraph.text.replace(find_text, replace_text)
                
                # Add track changes formatting (red strikethrough for deleted, blue underline for inserted)
                for run in paragraph.runs:
                    if replace_text in run.text:
                        run.font.color.rgb = RGBColor(0, 0, 255)  # Blue for insertions
                        run.underline = True
                
                found = True
                print(f"Successfully replaced '{find_text}' with '{replace_text}'")
                break
        
        if not found:
            print(f"Text '{find_text}' not found in document")
    
    elif instruction_type == 'delete':
        find_text = instruction['find']
        print(f"Looking to delete: '{find_text}'")
        
        found = False
        for paragraph in doc.paragraphs:
            if find_text in paragraph.text:
                print(f"Found text to delete in paragraph: '{paragraph.text[:100]}...'")
                
                # Replace with strikethrough formatting (simulating deletion)
                paragraph.text = paragraph.text.replace(find_text, find_text)
                for run in paragraph.runs:
                    if find_text in run.text:
                        run.font.color.rgb = RGBColor(255, 0, 0)  # Red for deletions
                        run.font.strike = True
                
                found = True
                print(f"Successfully marked '{find_text}' for deletion")
                break
        
        if not found:
            print(f"Text '{find_text}' not found in document")
    
    elif instruction_type == 'add':
        text_to_add = instruction['text']
        location = instruction.get('location', 'at the end')
        
        print(f"Adding text: '{text_to_add}' {location}")
        
        # Add at the end for simplicity
        new_paragraph = doc.add_paragraph()
        run = new_paragraph.add_run(text_to_add)
        run.font.color.rgb = RGBColor(0, 0, 255)  # Blue for insertions
        run.underline = True
        
        print(f"Successfully added text: '{text_to_add}'")

def parse_redline_instructions(instructions_text):
    """Parse redline instructions into structured format."""
    lines = [line.strip() for line in instructions_text.split('\n') if line.strip()]
    parsed = []
    
    for line in lines:
        instruction = parse_instruction(line)
        if instruction:
            parsed.append(instruction)
        else:
            print(f"⚠️  Could not parse instruction: '{line}'")
    
    return parsed

def parse_instruction(line):
    """Parse a single instruction line."""
    original = line
    
    # Pattern: Change "old" to "new"
    change_pattern1 = re.compile(r'^change\s+"([^"]+)"\s+to\s+"([^"]+)"(?:\s+(.+))?$', re.IGNORECASE)
    match = change_pattern1.match(line)
    if match:
        return {
            'type': 'replace',
            'original': original,
            'find': match.group(1),
            'replace': match.group(2),
            'context': match.group(3) if match.group(3) else None
        }
    
    # Pattern: Change [text] to [text] (without quotes)
    change_pattern2 = re.compile(r'^change\s+(.+?)\s+to\s+(.+?)$', re.IGNORECASE)
    match = change_pattern2.match(line)
    if match and '"' not in line:
        return {
            'type': 'replace',
            'original': original,
            'find': match.group(1).strip(),
            'replace': match.group(2).strip(),
            'context': None
        }
    
    # Pattern: Section X change - "old" to "new"
    section_pattern = re.compile(r'^section\s+\d+\s+change\s*-\s*"([^"]+)"\s+to\s+"([^"]+)"$', re.IGNORECASE)
    match = section_pattern.match(line)
    if match:
        return {
            'type': 'replace',
            'original': original,
            'find': match.group(1),
            'replace': match.group(2),
            'context': None
        }
    
    # Pattern: Replace "old" with "new"
    replace_pattern = re.compile(r'^replace\s+"([^"]+)"\s+with\s+"([^"]+)"(?:\s+(.+))?$', re.IGNORECASE)
    match = replace_pattern.match(line)
    if match:
        return {
            'type': 'replace',
            'original': original,
            'find': match.group(1),
            'replace': match.group(2),
            'context': match.group(3) if match.group(3) else None
        }
    
    # Pattern: Delete "text"
    delete_pattern = re.compile(r'^delete\s+"([^"]+)"(?:\s+(.+))?$', re.IGNORECASE)
    match = delete_pattern.match(line)
    if match:
        return {
            'type': 'delete',
            'original': original,
            'find': match.group(1),
            'context': match.group(2) if match.group(2) else None
        }
    
    # Pattern: Add "text"
    add_pattern = re.compile(r'^add\s+"([^"]+)"(?:\s+(.+))?$', re.IGNORECASE)
    match = add_pattern.match(line)
    if match:
        return {
            'type': 'add',
            'original': original,
            'text': match.group(1),
            'location': match.group(2) if match.group(2) else 'at the end'
        }
    
    return None

def process_word_document(input_path, instructions_text, output_path):
    """Process a Word document with redline instructions."""
    try:
        print(f"Loading document: {input_path}")
        doc = Document(input_path)
        
        print(f"Document loaded successfully. Paragraphs: {len(doc.paragraphs)}")
        
        # Parse instructions
        instructions = parse_redline_instructions(instructions_text)
        print(f"Parsed {len(instructions)} instructions")
        
        # Apply each instruction
        for instruction in instructions:
            apply_redline_instruction(doc, instruction)
        
        # Save the document
        print(f"Saving document to: {output_path}")
        doc.save(output_path)
        print("Document processing completed successfully!")
        
        return {
            'success': True,
            'message': f'Document processed successfully. Applied {len(instructions)} instructions.',
            'output_path': output_path,
            'instructions_processed': len(instructions)
        }
        
    except Exception as e:
        error_msg = f"Error processing document: {str(e)}"
        print(error_msg)
        return {
            'success': False,
            'error': error_msg
        }

def main():
    """Main entry point for the script."""
    parser = argparse.ArgumentParser(description='Process Word documents with redline instructions')
    parser.add_argument('input_file', help='Input Word document path')
    parser.add_argument('output_file', help='Output Word document path')
    parser.add_argument('--instructions', help='Redline instructions as text')
    parser.add_argument('--instructions-file', help='File containing redline instructions')
    
    args = parser.parse_args()
    
    # Get instructions
    if args.instructions:
        instructions_text = args.instructions
    elif args.instructions_file:
        with open(args.instructions_file, 'r') as f:
            instructions_text = f.read()
    else:
        print("Error: Either --instructions or --instructions-file must be provided")
        sys.exit(1)
    
    # Process the document
    result = process_word_document(args.input_file, instructions_text, args.output_file)
    
    # Output result as JSON
    print(json.dumps(result, indent=2))
    
    if not result['success']:
        sys.exit(1)

if __name__ == '__main__':
    main()