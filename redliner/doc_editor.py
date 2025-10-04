#!/usr/bin/env python3
"""Lightweight Word clause editor using python-docx.

This produces an updated copy of the document so you can use Word's
*Review ▸ Compare* feature to generate tracked changes in seconds.
"""
from __future__ import annotations

import argparse
import os
import re
import subprocess
import sys
from copy import deepcopy
from dataclasses import dataclass
from typing import List, Optional

try:
    from docx import Document
except ImportError as exc:
    raise SystemExit(
        "python-docx is required. Install with `python3 -m pip install python-docx`."
    ) from exc


INSTRUCTION_HEADER = re.compile(r"^\s*\[(?P<action>[A-Za-z_\- ]+)\]\s*$")
FIELD_PATTERN = re.compile(r"^\s*([A-Za-z_ ]+):\s*(.*)\s*$")


@dataclass
class Command:
    index: int
    action: str
    section: Optional[str]
    anchor: Optional[str]
    target: Optional[str]
    replacement: Optional[str]
    text: Optional[str]

    def label(self) -> str:
        bits = [f"#{self.index}", self.action.upper()]
        if self.section:
            bits.append(f"[{self.section}]")
        return " ".join(bits)


class InstructionParseError(Exception):
    """Raised when instruction text cannot be parsed."""


def read_clipboard() -> str:
    try:
        result = subprocess.run(["pbpaste"], capture_output=True, check=True, text=True)
    except (FileNotFoundError, subprocess.CalledProcessError):
        raise InstructionParseError(
            "Could not read clipboard. Copy your instructions and try again."
        )
    return result.stdout.strip()


def parse_instructions(raw: str) -> List[Command]:
    commands: List[Command] = []
    current: Optional[dict] = None
    waiting_field: Optional[str] = None
    buffer: List[str] = []
    lines = raw.splitlines()

    def flush_buffer() -> str:
        nonlocal buffer
        text = "\n".join(buffer).strip()
        buffer = []
        return text

    def finalize_current() -> None:
        nonlocal current
        if current is None:
            return
        action = current.get("action")
        if not action:
            raise InstructionParseError("Instruction missing header like [REPLACE].")
        cmd = Command(
            index=len(commands) + 1,
            action=action.strip().lower().replace(" ", "_"),
            section=current.get("section"),
            anchor=current.get("anchor"),
            target=current.get("target"),
            replacement=current.get("replacement"),
            text=current.get("text"),
        )
        commands.append(cmd)
        current = None

    for line in lines + ["---END---"]:
        header_match = INSTRUCTION_HEADER.match(line)
        if header_match:
            finalize_current()
            current = {"action": header_match.group("action")}
            waiting_field = None
            buffer = []
            continue
        if waiting_field and line.strip() == "---":
            current[waiting_field] = flush_buffer()
            waiting_field = None
            continue
        if line == "---END---":
            if waiting_field:
                current[waiting_field] = flush_buffer()
                waiting_field = None
            finalize_current()
            continue
        if waiting_field:
            buffer.append(line)
            continue
        field_match = FIELD_PATTERN.match(line)
        if field_match:
            field = field_match.group(1).strip().lower().replace(" ", "_")
            value = field_match.group(2).strip()
            if value == "<<<":
                waiting_field = field
                buffer = []
            else:
                if current is None:
                    raise InstructionParseError(
                        f"Field '{field}' found before any instruction header."
                    )
                current[field] = value
            continue
        if line.strip() == "":
            continue
        raise InstructionParseError(f"Unrecognized line: '{line}'")

    if not commands:
        raise InstructionParseError("No instructions found.")
    return commands


# --- Document helpers ------------------------------------------------------------

def find_and_replace(paragraphs: List[str], target: str, replacement: str) -> bool:
    for idx, text in enumerate(paragraphs):
        if target in text:
            paragraphs[idx] = text.replace(target, replacement, 1)
            return True
    return False


def insert_text(paragraphs: List[str], anchor: str, new_text: str, *, before: bool) -> bool:
    for idx, text in enumerate(paragraphs):
        if anchor in text:
            if before:
                paragraphs[idx] = text.replace(anchor, f"{new_text}\n{anchor}", 1)
            else:
                paragraphs[idx] = text.replace(anchor, f"{anchor}\n{new_text}", 1)
            return True
    return False


def delete_text(paragraphs: List[str], target: str) -> bool:
    return find_and_replace(paragraphs, target, "")


def load_paragraphs(document: Document) -> List[str]:
    return [para.text for para in document.paragraphs]


def write_paragraphs(document: Document, paragraphs: List[str]) -> None:
    for para, text in zip(document.paragraphs, paragraphs):
        para.text = text
    if len(document.paragraphs) < len(paragraphs):
        # Append extras if inserts created new lines beyond existing paragraph count
        for text in paragraphs[len(document.paragraphs) :]:
            document.add_paragraph(text)


def apply_command(paragraphs: List[str], command: Command) -> bool:
    if command.action == "replace" and command.target and command.replacement is not None:
        return find_and_replace(paragraphs, command.target, command.replacement)
    if command.action == "insert_after" and command.anchor and command.text:
        return insert_text(paragraphs, command.anchor, command.text, before=False)
    if command.action == "insert_before" and command.anchor and command.text:
        return insert_text(paragraphs, command.anchor, command.text, before=True)
    if command.action in {"delete", "remove"} and command.target:
        return delete_text(paragraphs, command.target)
    raise InstructionParseError(f"Unsupported or incomplete command: {command.label()}")


# --- CLI ------------------------------------------------------------------------

def main(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(description="Apply instruction-driven edits to a Word doc (no Track Changes).")
    parser.add_argument("document", help="Source .docx file")
    parser.add_argument(
        "--output",
        help="Write changes to this path (default adds _edited suffix).",
    )
    parser.add_argument(
        "--instructions",
        help="Optional path to a text file of instructions. Defaults to clipboard.",
    )
    args = parser.parse_args(argv)

    doc_path = os.path.abspath(args.document)
    if not os.path.exists(doc_path):
        print(f"Document not found: {doc_path}", file=sys.stderr)
        return 1

    if args.instructions:
        with open(args.instructions, "r", encoding="utf-8") as fh:
            raw_instructions = fh.read()
    else:
        raw_instructions = read_clipboard()

    try:
        commands = parse_instructions(raw_instructions)
    except InstructionParseError as exc:
        print(f"Instruction error: {exc}", file=sys.stderr)
        return 1

    document = Document(doc_path)
    paragraphs = load_paragraphs(document)
    updated = deepcopy(paragraphs)
    failures = 0

    for command in commands:
        ok = apply_command(updated, command)
        if ok:
            print(f"- {command.label()}: applied")
        else:
            print(f"- {command.label()}: target not found")
            failures += 1

    write_paragraphs(document, updated)

    output_path = args.output or os.path.splitext(doc_path)[0] + "_edited.docx"
    document.save(output_path)
    print(f"\nSaved updated copy to {output_path}")
    if failures:
        print("Some instructions were not matched. Review the log above.")
    print("Next: open both files in Word and use Review ▸ Compare to generate tracked redlines.")
    return 1 if failures else 0


if __name__ == "__main__":
    sys.exit(main())
