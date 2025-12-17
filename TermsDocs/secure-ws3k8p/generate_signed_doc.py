#!/usr/bin/env python3
"""
Generate a signed version of the Word document with signature embedded.
Uses Pillow to create the signature image.
"""

import subprocess
import sys
from pathlib import Path
from datetime import datetime
import io

# Install dependencies if needed
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx', '--user'])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError:
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'pillow', '--user'])
    from PIL import Image, ImageDraw, ImageFont

# Paths
WORKROOM_DIR = Path(__file__).parent
PNG_PATH = WORKROOM_DIR / 'signature.png'
ORIGINAL_DOC = WORKROOM_DIR / 'Demand Letter - Shaw v Sunnova - Rescission.docx'
SIGNED_DOC = WORKROOM_DIR / 'Demand Letter - Shaw v Sunnova - Rescission - SIGNED.docx'

def create_signature_image():
    """Create a handwritten-style signature image using Pillow."""
    print("Creating signature image...")

    # Create image with white background
    width, height = 400, 100
    img = Image.new('RGBA', (width, height), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)

    # Try to use a script font, fall back to default
    font_size = 38
    try:
        # Try macOS fonts
        for font_path in [
            '/System/Library/Fonts/Supplemental/Brush Script.ttf',
            '/System/Library/Fonts/Supplemental/Bradley Hand Bold.ttf',
            '/System/Library/Fonts/Supplemental/Snell Roundhand.ttf',
            '/Library/Fonts/Brush Script.ttf',
        ]:
            if Path(font_path).exists():
                font = ImageFont.truetype(font_path, font_size)
                break
        else:
            font = ImageFont.load_default()
    except Exception:
        font = ImageFont.load_default()

    # Draw signature text with dark blue ink color
    ink_color = (26, 54, 93, 230)  # #1a365d with some transparency

    # Draw the signature with Esq.
    draw.text((20, 20), "/s/ Sergei Tokmakov, Esq.", fill=ink_color, font=font)

    # Add a flourish underline
    draw.line([(15, 70), (380, 70)], fill=(26, 54, 93, 150), width=2)

    # Save as PNG
    img.save(str(PNG_PATH), 'PNG')
    print(f"Created: {PNG_PATH}")

def generate_verification_code():
    """Generate a UUID-like verification code."""
    import random
    chars = 'ABCDEF0123456789'
    segments = [8, 4, 4, 4, 12]
    return '-'.join(
        ''.join(random.choice(chars) for _ in range(length))
        for length in segments
    )

def create_signed_document():
    """Create signed version of the Word document."""
    print(f"Opening: {ORIGINAL_DOC}")
    doc = Document(ORIGINAL_DOC)

    # Generate verification code
    verification_code = generate_verification_code()
    timestamp = datetime.now().strftime('%B %d, %Y at %I:%M %p')

    # Find "Sincerely," and insert signature after it
    for i, para in enumerate(doc.paragraphs):
        if 'Sincerely,' in para.text:
            # Find the next paragraph(s) after "Sincerely,"
            for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                next_para = doc.paragraphs[j]
                text = next_para.text.strip()

                # If we hit the name line, insert signature before it
                if 'Sergei Tokmakov' in text:
                    # Insert signature block before the name
                    sig_para = next_para.insert_paragraph_before('')

                    # Add signature image
                    run = sig_para.add_run()
                    run.add_picture(str(PNG_PATH), width=Inches(3.0))

                    # Add signature details block
                    details_para = next_para.insert_paragraph_before('')

                    # Signed by line
                    run1 = details_para.add_run('✓ Signed by: Sergei Tokmakov, Esq.\n')
                    run1.font.size = Pt(9)
                    run1.font.color.rgb = RGBColor(22, 101, 52)  # Green

                    # CBN line
                    run2 = details_para.add_run('   CBN 279869\n')
                    run2.font.size = Pt(9)
                    run2.font.color.rgb = RGBColor(71, 85, 105)

                    # Date line
                    run3 = details_para.add_run(f'{timestamp}\n')
                    run3.font.size = Pt(9)
                    run3.font.color.rgb = RGBColor(71, 85, 105)

                    # Email line
                    run4 = details_para.add_run('Verified email: owner@terms.law\n')
                    run4.font.size = Pt(9)
                    run4.font.color.rgb = RGBColor(71, 85, 105)

                    # Verification code
                    run5 = details_para.add_run(verification_code)
                    run5.font.size = Pt(8)
                    run5.font.color.rgb = RGBColor(148, 163, 184)
                    run5.font.name = 'Courier New'

                    print(f"Signature inserted with verification code: {verification_code}")
                    break
            break

    # Save the signed document
    doc.save(SIGNED_DOC)
    print(f"Saved signed document: {SIGNED_DOC}")

def main():
    """Main function."""
    print("=" * 50)
    print("Generating Signed Word Document")
    print("=" * 50)

    # Check if original document exists
    if not ORIGINAL_DOC.exists():
        print(f"ERROR: Original document not found: {ORIGINAL_DOC}")
        sys.exit(1)

    # Create signature image
    create_signature_image()

    # Create signed document
    create_signed_document()

    print("\nDone!")

if __name__ == '__main__':
    main()
